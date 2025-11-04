"""
Exporter Service - Convert generated markdown sections to Word documents.

This service handles the conversion of markdown content to professional Word documents
with Bernadet styling (Arial fonts, custom colors).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
from typing import List, Dict
import os


class ExporterService:
    """Service for exporting generated sections to Word documents."""

    def __init__(self, template_path: str = "./templates/template.docx"):
        """
        Initialize the exporter service.

        Args:
            template_path: Path to Word template file (optional)
        """
        self.template_path = template_path

    def create_memoire(
        self,
        project_name: str,
        sections: List[Dict]
    ) -> str:
        """
        Create a Word document from generated sections.

        Args:
            project_name: Name of the project
            sections: List of sections with 'title' and 'content' (markdown)
                     Expected format: [{'title': str, 'content': str}, ...]

        Returns:
            Path to generated .docx file
        """
        # Try to load template, otherwise create new document
        try:
            if os.path.exists(self.template_path):
                doc = Document(self.template_path)
            else:
                doc = Document()
                self._apply_default_styles(doc)
        except Exception:
            # If template fails to load, create new document
            doc = Document()
            self._apply_default_styles(doc)

        # Add title page
        doc.add_heading(f"Mémoire Technique - {project_name}", level=0)
        doc.add_page_break()

        # Add each section
        for section in sections:
            title = section.get('title', 'Section')
            content = section.get('content', '')

            # Add section
            self._add_section(doc, title, content)
            doc.add_page_break()

        # Ensure data directory exists
        os.makedirs('./data', exist_ok=True)

        # Save document
        safe_project_name = project_name.replace(' ', '_').replace('/', '_')
        output_path = f"./data/output_{safe_project_name}.docx"
        doc.save(output_path)

        return output_path

    def _add_section(self, doc: Document, title: str, markdown_content: str):
        """
        Add a section to the document with markdown conversion.

        Args:
            doc: python-docx Document object
            title: Section title
            markdown_content: Section content in markdown format
        """
        # Add section title
        doc.add_heading(title, level=1)

        # Convert markdown to HTML
        html = markdown.markdown(markdown_content, extensions=['tables'])
        soup = BeautifulSoup(html, 'html.parser')

        # Parse HTML and add to document
        for element in soup.find_all(['h2', 'h3', 'p', 'ul', 'ol', 'table']):
            if element.name == 'h2':
                doc.add_heading(element.text, level=2)

            elif element.name == 'h3':
                doc.add_heading(element.text, level=3)

            elif element.name == 'p':
                # Skip empty paragraphs
                if element.text.strip():
                    doc.add_paragraph(element.text)

            elif element.name == 'ul':
                # Bullet list
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.text.strip(), style='List Bullet')

            elif element.name == 'ol':
                # Numbered list
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.text.strip(), style='List Number')

            elif element.name == 'table':
                self._add_table(doc, element)

    def _add_table(self, doc: Document, html_table):
        """
        Convert HTML table to Word table.

        Args:
            doc: python-docx Document object
            html_table: BeautifulSoup table element
        """
        rows = html_table.find_all('tr')
        if not rows:
            return

        # Count columns from first row
        first_row_cells = rows[0].find_all(['th', 'td'])
        if not first_row_cells:
            return

        cols = len(first_row_cells)

        # Create Word table
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < cols:  # Safety check
                    table.rows[i].cells[j].text = cell.text.strip()

        # Add spacing after table
        doc.add_paragraph()

    def _apply_default_styles(self, doc: Document):
        """
        Apply Bernadet default styles to the document.

        Bernadet Brand Colors:
        - Heading 1: Blue RGB(46, 80, 144)
        - Heading 2: Green RGB(120, 180, 90)
        - Font: Arial throughout
        """
        try:
            # Style Heading 1
            h1 = doc.styles['Heading 1']
            h1.font.name = 'Arial'
            h1.font.size = Pt(18)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(46, 80, 144)  # Blue

            # Style Heading 2
            h2 = doc.styles['Heading 2']
            h2.font.name = 'Arial'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(120, 180, 90)  # Green

            # Style Heading 3
            h3 = doc.styles['Heading 3']
            h3.font.name = 'Arial'
            h3.font.size = Pt(12)
            h3.font.bold = True

            # Style Normal text
            normal = doc.styles['Normal']
            normal.font.name = 'Arial'
            normal.font.size = Pt(11)
        except KeyError:
            # If style doesn't exist, skip
            pass
