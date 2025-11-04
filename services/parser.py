"""
Parser Service for extracting text from PDF and DOCX files.
"""
import os
import re
from typing import List, Dict, Any, Set
from pypdf import PdfReader
from docx import Document
from models.schemas import ParsedSection, ParseResult, ChunkData


class ParserService:
    """Service for parsing PDF and DOCX documents."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        """
        Initialize the parser service.

        Args:
            chunk_size: Number of characters per chunk
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text by removing noise and excessive whitespace.

        This method:
        - Removes excessive whitespace (multiple spaces, tabs, newlines)
        - Normalizes line breaks
        - Removes common document artifacts (page numbers, etc.)
        - Preserves meaningful paragraph structure

        Args:
            text: Raw text to clean

        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""

        # Remove form feed and other control characters (except newlines and tabs)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

        # Remove excessive whitespace while preserving paragraph breaks
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)

        # Replace tabs with spaces
        text = text.replace('\t', ' ')

        # Normalize line breaks: replace 3+ newlines with 2 newlines (paragraph break)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove spaces at the beginning and end of lines
        text = re.sub(r' *\n *', '\n', text)

        # Remove common page number patterns
        # Pattern: "Page X", "Page X of Y", "X / Y", "- X -", etc.
        text = re.sub(r'\n\s*-?\s*\d+\s*-?\s*\n', '\n', text)  # Standalone page numbers
        text = re.sub(r'\n\s*Page\s+\d+(\s+of\s+\d+)?\s*\n', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\s*\d+\s*/\s*\d+\s*\n', '\n', text)  # "X / Y" format

        # Remove common header/footer patterns (often repeated)
        # Pattern: Short lines at start/end that might be headers/footers
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            stripped = line.strip()
            # Skip very short lines that are likely artifacts (but keep empty lines for paragraphs)
            if len(stripped) == 0:
                cleaned_lines.append('')
            elif len(stripped) < 3:  # Skip lines with only 1-2 characters (likely artifacts)
                continue
            else:
                cleaned_lines.append(line)

        text = '\n'.join(cleaned_lines)

        # Final cleanup: remove any remaining multiple spaces
        text = re.sub(r' +', ' ', text)

        # Normalize excessive newlines again after line filtering
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text

    def detect_repeated_patterns(self, text: str, min_length: int = 10, max_length: int = 100) -> Set[str]:
        """
        Detect repeated patterns in text (likely headers/footers).

        Looks for strings that appear multiple times and might be headers/footers.

        Args:
            text: Text to analyze
            min_length: Minimum length of pattern to detect
            max_length: Maximum length of pattern to detect

        Returns:
            Set of repeated patterns found
        """
        patterns = set()
        lines = text.split('\n')

        # Count occurrences of each line
        line_counts = {}
        for line in lines:
            stripped = line.strip()
            # Only consider lines within length bounds
            if min_length <= len(stripped) <= max_length:
                line_counts[stripped] = line_counts.get(stripped, 0) + 1

        # Consider a line a "repeated pattern" if it appears 3+ times
        # (likely a header/footer if it appears on multiple pages)
        for line, count in line_counts.items():
            if count >= 3:
                patterns.add(line)

        return patterns

    def remove_repeated_patterns(self, text: str) -> str:
        """
        Remove repeated patterns (headers/footers) from text.

        Args:
            text: Text to clean

        Returns:
            Text with repeated patterns removed
        """
        # Detect repeated patterns
        patterns = self.detect_repeated_patterns(text)

        if not patterns:
            return text

        # Remove each pattern
        for pattern in patterns:
            # Use word boundaries to avoid removing parts of legitimate text
            text = text.replace(f'\n{pattern}\n', '\n')
            text = text.replace(f'{pattern}\n', '')
            text = text.replace(f'\n{pattern}', '')

        return text

    def parse_pdf(self, filepath: str) -> ParseResult:
        """
        Parse a PDF file and extract text.

        Args:
            filepath: Path to the PDF file

        Returns:
            ParseResult with extracted text and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is corrupted or unreadable
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        try:
            reader = PdfReader(filepath)
            sections = []

            # Extract text from each page
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():  # Only add non-empty pages
                    sections.append(ParsedSection(
                        title=f"Page {page_num}",
                        content=text,
                        page=page_num,
                        order=page_num
                    ))

            # Combine all text
            full_text = '\n\n'.join(section.content for section in sections)

            # Clean the text to remove whitespace, headers, footers, etc.
            full_text = self.clean_text(full_text)
            full_text = self.remove_repeated_patterns(full_text)

            # Extract metadata
            metadata = {}
            if reader.metadata:
                metadata = {
                    'author': reader.metadata.get('/Author', None),
                    'creator': reader.metadata.get('/Creator', None),
                    'producer': reader.metadata.get('/Producer', None),
                    'subject': reader.metadata.get('/Subject', None),
                    'title': reader.metadata.get('/Title', None),
                }

            return ParseResult(
                sections=sections,
                full_text=full_text,
                page_count=len(reader.pages),
                char_count=len(full_text),
                metadata=metadata
            )

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    def parse_docx(self, filepath: str) -> ParseResult:
        """
        Parse a DOCX file and extract text.

        Args:
            filepath: Path to the DOCX file

        Returns:
            ParseResult with extracted text and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is corrupted or unreadable
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        try:
            doc = Document(filepath)
            sections = []
            current_section = ParsedSection(
                title='Introduction',
                content='',
                order=0
            )
            section_order = 0

            # Extract text from paragraphs
            for para in doc.paragraphs:
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    # Save current section if it has content
                    if current_section.content.strip():
                        sections.append(current_section)
                        section_order += 1

                    # Start new section
                    current_section = ParsedSection(
                        title=para.text,
                        content='',
                        order=section_order
                    )
                else:
                    # Add paragraph to current section
                    if para.text.strip():
                        current_section.content += para.text + '\n'

            # Add the last section
            if current_section.content.strip():
                sections.append(current_section)

            # If no sections were created (no headings), create one section with all content
            if not sections:
                all_text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
                sections.append(ParsedSection(
                    title='Document',
                    content=all_text,
                    order=0
                ))

            # Combine all text
            full_text = '\n\n'.join(section.content for section in sections)

            # Clean the text to remove whitespace, headers, footers, etc.
            full_text = self.clean_text(full_text)
            full_text = self.remove_repeated_patterns(full_text)

            # Extract metadata from core properties
            metadata = {}
            if doc.core_properties:
                core_props = doc.core_properties
                metadata = {
                    'author': core_props.author,
                    'created': str(core_props.created) if core_props.created else None,
                    'modified': str(core_props.modified) if core_props.modified else None,
                    'title': core_props.title,
                    'subject': core_props.subject,
                }

            return ParseResult(
                sections=sections,
                full_text=full_text,
                paragraph_count=len(doc.paragraphs),
                char_count=len(full_text),
                metadata=metadata
            )

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[ChunkData]:
        """
        Split text into chunks for RAG processing.

        Uses a sliding window approach with overlap to maintain context.

        Args:
            text: The text to chunk
            metadata: Optional metadata to include with each chunk

        Returns:
            List of ChunkData objects
        """
        if metadata is None:
            metadata = {}

        # Clean text before chunking (if not already cleaned)
        text = self.clean_text(text)

        chunks = []
        text_length = len(text)
        chunk_index = 0

        # Handle empty or very short text
        if text_length == 0:
            return chunks

        if text_length <= self.chunk_size:
            # Text is shorter than chunk size, return as single chunk
            chunks.append(ChunkData(
                content=text.strip(),
                chunk_index=0,
                char_start=0,
                char_end=text_length,
                metadata=metadata
            ))
            return chunks

        # Sliding window chunking with overlap
        start = 0
        while start < text_length:
            end = min(start + self.chunk_size, text_length)
            chunk_text = text[start:end]

            # Only add non-empty chunks
            if chunk_text.strip():
                chunks.append(ChunkData(
                    content=chunk_text.strip(),
                    chunk_index=chunk_index,
                    char_start=start,
                    char_end=end,
                    metadata={
                        **metadata,
                        'total_chunks': None,  # Will be set after we know total
                    }
                ))
                chunk_index += 1

            # Move forward by (chunk_size - overlap)
            start += (self.chunk_size - self.chunk_overlap)

            # Prevent infinite loop on edge cases
            if start >= text_length:
                break

        # Update total_chunks in metadata
        for chunk in chunks:
            chunk.metadata['total_chunks'] = len(chunks)

        return chunks

    def parse_file(self, filepath: str) -> ParseResult:
        """
        Parse a file based on its extension.

        Args:
            filepath: Path to the file

        Returns:
            ParseResult with extracted text

        Raises:
            ValueError: If file type is not supported
        """
        ext = os.path.splitext(filepath)[1].lower()

        if ext == '.pdf':
            return self.parse_pdf(filepath)
        elif ext in ['.docx', '.doc']:
            return self.parse_docx(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Only .pdf and .docx are supported.")

    def get_file_metadata(self, filepath: str) -> Dict[str, Any]:
        """
        Extract basic file metadata.

        Args:
            filepath: Path to the file

        Returns:
            Dictionary with file metadata
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        stat = os.stat(filepath)
        return {
            'size_bytes': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'extension': os.path.splitext(filepath)[1].lower(),
            'filename': os.path.basename(filepath),
        }
